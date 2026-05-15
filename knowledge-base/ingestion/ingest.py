#!/usr/bin/env python3
"""
RAG ingestion pipeline: PDF/DOCX/TXT/MD -> chunks -> embeddings -> Qdrant.

Usage:
    python ingest.py --source ../raw/pharmacy --collection knowledge_base
"""

import argparse
import hashlib
import json
import logging
import re
import sys
import uuid
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import httpx
import time
import math
from docx import Document
from PyPDF2 import PdfReader
from tqdm import tqdm


SUPPORTED_EXTENSIONS = {".docx", ".md", ".txt", ".pdf"}
DEFAULT_SECTION = "General"
logger = logging.getLogger(__name__)


def is_valid_vector(vec: list[float]) -> bool:
    """Проверить, что вектор не содержит NaN или Inf."""
    if not vec:
        return False
    return not any(math.isnan(x) or math.isinf(x) for x in vec)



@dataclass(slots=True)
class Segment:
    section: str
    text: str


@dataclass(slots=True)
class LoadedDocument:
    source_file: str
    category: str
    file_type: str
    document_title: str
    sections: list[Segment]


@dataclass(slots=True)
class Chunk:
    id: str
    content: str
    source_file: str
    document_title: str
    section: str
    category: str
    file_type: str
    chunk_index: int
    total_chunks: int


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\u00a0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def infer_category(source: Path, root: Path) -> str:
    rel_parts = source.relative_to(root).parts
    return rel_parts[0] if rel_parts else "unknown"


def infer_title(source: Path) -> str:
    return source.stem.replace("_", " ").strip()


def parse_markdown_sections(text: str) -> list[Segment]:
    segments: list[Segment] = []
    current_section = DEFAULT_SECTION
    buffer: list[str] = []

    def flush() -> None:
        if not buffer:
            return
        chunk_text = normalize_text("\n".join(buffer))
        if chunk_text:
            logger.debug("Parsed markdown section '%s' with %d characters", current_section, len(chunk_text))
            segments.append(Segment(section=current_section, text=chunk_text))
        buffer.clear()

    for line in text.splitlines():
        heading_match = re.match(r"^#{1,6}\s+(.*)$", line.strip())
        if heading_match:
            flush()
            current_section = heading_match.group(1).strip() or DEFAULT_SECTION
            continue
        buffer.append(line)

    flush()
    if not segments:
        normalized = normalize_text(text)
        if normalized:
            segments.append(Segment(section=DEFAULT_SECTION, text=normalized))
    return segments


def parse_plain_sections(text: str) -> list[Segment]:
    normalized = normalize_text(text)
    if not normalized:
        return []
    logger.debug("Parsed plain text section with %d characters", len(normalized))
    return [Segment(section=DEFAULT_SECTION, text=normalized)]


def parse_docx_sections(path: Path) -> list[Segment]:
    logger.debug("Parsing DOCX document: %s", path)
    document = Document(str(path))
    segments: list[Segment] = []
    current_section = DEFAULT_SECTION
    buffer: list[str] = []

    def flush() -> None:
        if not buffer:
            return
        text_block = normalize_text("\n".join(buffer))
        if text_block:
            logger.debug("Parsed DOCX section '%s' with %d characters", current_section, len(text_block))
            segments.append(Segment(section=current_section, text=text_block))
        buffer.clear()

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = ""
        if paragraph.style is not None and paragraph.style.name is not None:
            style_name = paragraph.style.name
        style_name = style_name.lower()
        is_heading_style = style_name.startswith("heading")
        is_heading_text = bool(re.match(r"^(\d+(\.\d+)*|§)\s+", text))
        if is_heading_style or is_heading_text:
            flush()
            current_section = text
            continue

        buffer.append(text)

    for table in document.tables:
        for row in table.rows:
            row_values = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if row_values:
                buffer.append(" | ".join(row_values))

    flush()
    return segments


def parse_pdf_sections(path: Path) -> list[Segment]:
    logger.debug("Parsing PDF document: %s", path)
    reader = PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(page_text)
    return parse_plain_sections("\n\n".join(pages))


def load_document(path: Path, source_root: Path) -> LoadedDocument | None:
    ext = path.suffix.lower()
    raw_text = ""
    sections: list[Segment] = []
    logger.debug("Loading document: %s", path)

    if ext == ".md":
        raw_text = path.read_text(encoding="utf-8", errors="ignore")
        sections = parse_markdown_sections(raw_text)
    elif ext == ".txt":
        raw_text = path.read_text(encoding="utf-8", errors="ignore")
        sections = parse_plain_sections(raw_text)
    elif ext == ".docx":
        sections = parse_docx_sections(path)
    elif ext == ".pdf":
        sections = parse_pdf_sections(path)
    else:
        return None

    if not sections:
        logger.debug("Skipping document with no extractable sections: %s", path)
        return None

    source_file = str(path.relative_to(source_root)).replace("\\", "/")
    return LoadedDocument(
        source_file=source_file,
        category=infer_category(path, source_root),
        file_type=ext.lstrip("."),
        document_title=infer_title(path),
        sections=sections,
    )


def split_with_overlap(text: str, chunk_size: int, overlap: int) -> list[str]:
    if chunk_size <= 0:
        raise ValueError("chunk_size must be > 0")
    if overlap < 0:
        raise ValueError("overlap must be >= 0")
    if overlap >= chunk_size:
        raise ValueError("overlap must be less than chunk_size")

    normalized = normalize_text(text)
    if not normalized:
        return []

    chunks: list[str] = []
    start = 0
    length = len(normalized)

    while start < length:
        end = min(start + chunk_size, length)
        piece = normalized[start:end]
        if end < length:
            split_point = piece.rfind(" ")
            if split_point > int(chunk_size * 0.6):
                end = start + split_point
                piece = normalized[start:end]

        piece = piece.strip()
        if piece:
            chunks.append(piece)

        if end >= length:
            break
        start = max(end - overlap, start + 1)

    return chunks


def build_chunk_id(source_file: str, section: str, chunk_index: int) -> str:
    base = f"{source_file}::{section}::{chunk_index}"
    digest = hashlib.sha1(base.encode("utf-8")).hexdigest()
    return str(uuid.uuid5(uuid.NAMESPACE_URL, digest))


def build_chunks(documents: list[LoadedDocument], chunk_size: int, overlap: int) -> list[Chunk]:
    chunks: list[Chunk] = []

    for doc in documents:
        logger.debug("Building chunks for document '%s' (%s)", doc.document_title, doc.source_file)
        doc_chunks: list[Chunk] = []
        chunk_counter = 0
        for segment in doc.sections:
            split_parts = split_with_overlap(segment.text, chunk_size=chunk_size, overlap=overlap)
            for content in split_parts:
                chunk_counter += 1
                doc_chunks.append(
                    Chunk(
                        id=build_chunk_id(doc.source_file, segment.section, chunk_counter),
                        content=content,
                        source_file=doc.source_file,
                        document_title=doc.document_title,
                        section=segment.section,
                        category=doc.category,
                        file_type=doc.file_type,
                        chunk_index=chunk_counter,
                        total_chunks=0,
                    )
                )

        total = len(doc_chunks)
        for chunk in doc_chunks:
            chunk.total_chunks = total
        chunks.extend(doc_chunks)

    logger.debug("Built %d total chunks from %d documents", len(chunks), len(documents))
    return chunks


def ensure_collection(
    client: httpx.Client,
    qdrant_url: str,
    collection: str,
    vector_size: int,
    recreate: bool,
) -> None:
    base = qdrant_url.rstrip("/")
    collection_url = f"{base}/collections/{collection}"
    logger.info("Ensuring Qdrant collection '%s' at %s", collection, collection_url)

    if recreate:
        logger.warning("Recreating Qdrant collection '%s'", collection)
        client.delete(collection_url, timeout=30.0)

    response = client.get(collection_url, timeout=30.0)
    if response.status_code == 200:
        return
    if response.status_code not in (404,):
        raise RuntimeError(f"Failed to check Qdrant collection: {response.status_code} {response.text}")

    payload = {
        "vectors": {
            "size": vector_size,
            "distance": "Cosine",
        }
    }
    create_response = client.put(collection_url, json=payload, timeout=30.0)
    if create_response.status_code not in (200, 201):
        raise RuntimeError(
            f"Failed to create collection '{collection}': "
            f"{create_response.status_code} {create_response.text}"
        )
    logger.info("Qdrant collection '%s' created with vector size %d", collection, vector_size)


def embed_text(client: httpx.Client, ollama_url: str, model: str, text: str) -> list[float]:
    endpoint = f"{ollama_url.rstrip('/')}/api/embeddings"
    payload = {
        "model": model,
        "prompt": text,
    }
    # Retry transient errors (502/503/504) a few times with backoff so ingestion
    # can tolerate model cold-starts or temporary runner failures.
    max_attempts = 4
    backoff_base = 1.5
    for attempt in range(1, max_attempts + 1):
        logger.debug(
            "Embedding request attempt %d/%d: model=%s, text_len=%d, endpoint=%s",
            attempt,
            max_attempts,
            model,
            len(text),
            endpoint,
        )
        try:
            response = client.post(endpoint, json=payload, timeout=120.0)
        except Exception as ex:
            # Network/connection error
            if attempt == max_attempts:
                logger.exception("Embedding request error after %d attempts", max_attempts)
                raise RuntimeError(f"Embedding request error: {ex}") from ex
            wait = backoff_base ** attempt
            logger.warning(
                "Embedding request exception (attempt %d/%d): %s; retrying in %.1fs",
                attempt,
                max_attempts,
                ex,
                wait,
            )
            time.sleep(wait)
            continue

        if response.status_code == 200:
            data = response.json()
            embedding = data.get("embedding", [])
            if not isinstance(embedding, list) or not embedding:
                raise RuntimeError("Embedding response does not contain a valid vector")
            if not is_valid_vector(embedding):
                raise RuntimeError("Embedding contains NaN or Inf values")
            logger.debug("Received embedding vector with %d dimensions", len(embedding))
            return embedding

        # For transient server errors, retry; otherwise raise with body for debugging
        if response.status_code in (502, 503, 504):
            if attempt == max_attempts:
                logger.error(
                    "Embedding request failed after %d attempts: %s %s",
                    max_attempts,
                    response.status_code,
                    response.text,
                )
                raise RuntimeError(f"Embedding request failed: {response.status_code} {response.text}")
            wait = backoff_base ** attempt
            logger.warning(
                "Embedding returned %s (attempt %d/%d), retrying in %.1fs",
                response.status_code,
                attempt,
                max_attempts,
                wait,
            )
            time.sleep(wait)
            continue

        # Non-retriable error
        logger.error("Embedding request failed: %s %s", response.status_code, response.text)
        raise RuntimeError(f"Embedding request failed: {response.status_code} {response.text}")

    # Unreachable, but required for type checker (guarantee all paths return or raise)
    raise RuntimeError("Embedding request exhausted all retry attempts")


def to_qdrant_point(chunk: Chunk, vector: list[float], now_iso: str) -> dict[str, Any]:
    payload = {
        "content": chunk.content,
        "source_file": chunk.source_file,
        "document_title": chunk.document_title,
        "section": chunk.section,
        "category": chunk.category,
        "file_type": chunk.file_type,
        "chunk_index": str(chunk.chunk_index),
        "total_chunks": str(chunk.total_chunks),
        "language": "ru",
        "updated_at": now_iso,
        "version": "v1",
    }
    return {
        "id": chunk.id,
        "vector": vector,
        "payload": payload,
    }


def upsert_points(
    client: httpx.Client,
    qdrant_url: str,
    collection: str,
    points: list[dict[str, Any]],
    batch_size: int,
) -> None:
    endpoint = f"{qdrant_url.rstrip('/')}/collections/{collection}/points?wait=true"
    for i in range(0, len(points), batch_size):
        batch = points[i : i + batch_size]
        payload = {"points": batch}
        logger.debug(
            "Upserting batch %d-%d/%d into Qdrant collection '%s'",
            i + 1,
            i + len(batch),
            len(points),
            collection,
        )
        response = client.put(endpoint, json=payload, timeout=120.0)
        if response.status_code not in (200, 201):
            raise RuntimeError(f"Qdrant upsert failed: {response.status_code} {response.text}")


def discover_files(source_root: Path) -> list[Path]:
    files: list[Path] = []
    for path in source_root.rglob("*"):
        if not path.is_file():
            continue
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        files.append(path)
    files.sort()
    logger.debug("Discovered %d supported files under %s", len(files), source_root)
    return files


def main() -> None:
    parser = argparse.ArgumentParser(description="Ingest knowledge base documents into Qdrant")
    parser.add_argument("--source", default="../raw/pharmacy", help="Path to source documents directory")
    parser.add_argument("--collection", default="knowledge_base", help="Qdrant collection name")
    parser.add_argument("--ollama-url", default="http://localhost:11434")
    parser.add_argument("--qdrant-url", default="http://localhost:6333")
    parser.add_argument("--embedding-model", default="bge-m3")
    parser.add_argument("--chunk-size", type=int, default=600)
    parser.add_argument("--chunk-overlap", type=int, default=50)
    parser.add_argument("--batch-size", type=int, default=1)
    parser.add_argument("--log-level", default="INFO", help="Logging level (DEBUG, INFO, WARNING, ERROR)")
    parser.add_argument("--recreate", action="store_true", help="Recreate collection before upload")
    parser.add_argument("--dry-run", action="store_true", help="Process files without uploading to Qdrant")
    args = parser.parse_args()

    logging.basicConfig(
        level=getattr(logging, args.log_level.upper(), logging.INFO),
        format="%(asctime)s [%(levelname)s] %(message)s",
    )
    logger.info("Starting ingestion pipeline")
    logger.info(
        "Configuration: source=%s collection=%s ollama_url=%s qdrant_url=%s embedding_model=%s chunk_size=%d overlap=%d batch_size=%d recreate=%s dry_run=%s",
        args.source,
        args.collection,
        args.ollama_url,
        args.qdrant_url,
        args.embedding_model,
        args.chunk_size,
        args.chunk_overlap,
        args.batch_size,
        args.recreate,
        args.dry_run,
    )

    source_root = Path(args.source).resolve()
    if not source_root.exists() or not source_root.is_dir():
        logger.error("Source directory does not exist: %s", source_root)
        sys.exit(1)

    discovered = discover_files(source_root)
    if not discovered:
        logger.warning("No supported files found in: %s", source_root)
        logger.warning("Supported extensions: %s", ", ".join(sorted(SUPPORTED_EXTENSIONS)))
        sys.exit(0)

    documents: list[LoadedDocument] = []
    skipped_no_text = 0
    failed_files: list[tuple[Path, str]] = []

    print(f"[INFO] Source: {source_root}")
    print(f"[INFO] Files discovered: {len(discovered)}")

    for file_path in tqdm(discovered, desc="Loading documents", unit="file"):
        try:
            loaded = load_document(file_path, source_root)
            if loaded is None:
                skipped_no_text += 1
                continue
            documents.append(loaded)
        except Exception as ex:  # noqa: BLE001
            logger.exception("Failed to load document: %s", file_path)
            failed_files.append((file_path, str(ex)))

    chunks = build_chunks(documents, chunk_size=args.chunk_size, overlap=args.chunk_overlap)
    if not chunks:
        logger.warning("No chunks generated. Nothing to ingest.")
        if failed_files:
            logger.info("Failed files:")
            for path, reason in failed_files:
                logger.info("  - %s: %s", path, reason)
        sys.exit(0)

    logger.info("Documents loaded: %d", len(documents))
    logger.info("Chunks generated: %d", len(chunks))
    logger.info("Skipped (no text): %d", skipped_no_text)
    logger.info("Failed files: %d", len(failed_files))

    if args.dry_run:
        preview = {
            "chunk_id": chunks[0].id,
            "source_file": chunks[0].source_file,
            "section": chunks[0].section,
            "content_preview": chunks[0].content[:220],
        }
        logger.info("DRY RUN first chunk preview:\n%s", json.dumps(preview, ensure_ascii=False, indent=2))
        sys.exit(0)

    now_iso = datetime.now(timezone.utc).isoformat()

    with httpx.Client() as client:
        points: list[dict[str, Any]] = []
        # Probe call для определения размера вектора ДО основного loop
        logger.info("Running probe call to determine vector dimensions...")
        try:
            probe_vector = embed_text(client, args.ollama_url, args.embedding_model, "test")
            vector_size = len(probe_vector)
            logger.info("Vector size determined: %d dimensions", vector_size)
            ensure_collection(
                client=client,
                qdrant_url=args.qdrant_url,
                collection=args.collection,
                vector_size=vector_size,
                recreate=args.recreate,
            )
        except Exception as e:
            logger.error("Probe call failed: %s", e)
            raise RuntimeError("Cannot determine embedding dimensions") from e

        skipped_chunks = 0

        for chunk in tqdm(chunks, desc="Embedding chunks", unit="chunk"):
            logger.debug(
                "Embedding chunk %d/%d from %s [%s]",
                chunk.chunk_index,
                chunk.total_chunks,
                chunk.source_file,
                chunk.section,
            )
            try:
                vector = embed_text(client, args.ollama_url, args.embedding_model, chunk.content)
                if len(vector) != vector_size:
                    logger.error(
                        "Vector size mismatch for %s: expected %d, got %d",
                        chunk.source_file,
                        vector_size,
                        len(vector),
                    )
                    skipped_chunks += 1
                    continue
                points.append(to_qdrant_point(chunk, vector, now_iso))
            except Exception as e:
                logger.error("Skipping bad chunk %s: %s; content_preview: %s", chunk.id, e, chunk.content[:300])
                skipped_chunks += 1
                continue

        if not points:
            logger.error("No valid points to upsert. Ingestion failed.")
            sys.exit(1)

        logger.info("Upserting %d points (%d chunks skipped)", len(points), skipped_chunks)

        upsert_points(
            client=client,
            qdrant_url=args.qdrant_url,
            collection=args.collection,
            points=points,
            batch_size=args.batch_size,
        )

    logger.info("Ingestion completed. Uploaded points: %d", len(chunks))
    if failed_files:
        logger.info("Failed files:")
        for path, reason in failed_files:
            logger.info("  - %s: %s", path, reason)


if __name__ == "__main__":
    main()
